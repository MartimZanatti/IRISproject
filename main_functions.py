import os
import rouge
from  doc_class import doc_class, create_paragraph_stanza
from pre_processing import  find_final_data, is_title, before_italic_and_small_paragraphs
from embeddings import create_sim_matrix_word_2_vec, create_sim_matrix_bert, join_all_stanza_sentences
from lex_rank import degree_centrality_scores
from pos_processing import create_dict_scores, get_paragraphs, get_n_best_paragraphs
from docx import *
from rouge_evaluator import stanza_to_text
import numpy as np
import matplotlib.pyplot as plt


import nltk
#nltk.download('punkt')

def process_docx_file(file_name, preprocessing= True):

    d = doc_class(file_name, preprocessing)
    # pre-processamento para limitar as frases sumarizaveis
    if preprocessing:
        find_final_data(d.paragraphs)
        is_title(d.paragraphs)
        before_italic_and_small_paragraphs(d.paragraphs)

    return d

def create_stanza_sentences(doc):
    for p in doc.paragraphs:
        if p.sumarizable == True:
            nlp_paragraph = create_paragraph_stanza(p.text.get_text())
            p.stanza_text = nlp_paragraph
    return doc

def summarization(doc):

    similarity_matrix, ids_dict = create_sim_matrix_word_2_vec(doc.paragraphs, bert=False)
    #similarity_matrix, ids_dict = create_sim_matrix_bert(doc.paragraphs)

    scores = degree_centrality_scores(similarity_matrix.numpy(), threshold=None)

    #print(ids_dict)
    return scores, ids_dict


def pos_processing_paragraphs(paragraphs, scores, ids_dict):

    paragraphs_scores = create_dict_scores(scores, ids_dict)
    final_paragraphs = get_paragraphs(paragraphs, paragraphs_scores)

    return final_paragraphs

path = '../IrisDataset/automatic_sumaries/20210309_1sec_word_2_vec_no_pre_processing/'


def process_sum_to_files(paragraphs, scores, ids_dict, file_name, num_paragraphs):


    paragraphs_scores = create_dict_scores(scores, ids_dict)
    final_paragraphs = get_n_best_paragraphs(paragraphs, paragraphs_scores, num_paragraphs)

    documet = Document()

    for p in final_paragraphs:
        documet.add_paragraph(p)

    documet.save(path + file_name)



def rouge_main(data, real_sum_data, rouge_path):
    files = os.listdir(data)
    for file_name in files:
        if "txt" not in file_name:
            print(file_name)
            doc = doc_class(data + file_name, True)
            doc = create_stanza_sentences(doc)
            stanza_text = stanza_to_text(doc.paragraphs)
            real_sum = doc_class(real_sum_data + file_name, True)
            real_sum = create_stanza_sentences(real_sum)
            stanza_text_real_sum = stanza_to_text(real_sum.paragraphs)
            evaluator = rouge.Rouge(metrics=['rouge-n', 'rouge-l', "rouge-w"], max_n=4, limit_length=False,
                                    length_limit=100,
                                    length_limit_type='words',
                                    alpha=0.5,  # Default F1_score
                                    weight_factor=1.2,
                                    stemming=False)
            scores = evaluator.get_scores(stanza_text, stanza_text_real_sum)
            file = open(rouge_path + file_name[0:-6] + ".txt", "w", encoding="UTF-8")
            for metric, results in sorted(scores.items(), key=lambda x: x[0]):
                file.write(metric + "\n")
                file.write("f: ")
                file.write(str(results["f"]))
                file.write("\n")
                file.write("p: ")
                file.write(str(results["p"]))
                file.write("\n")
                file.write("r: ")
                file.write(str(results["r"]))
                file.write("\n")
            file.close()


def box_plot_main(paths):
    rouge_1_f = {}
    rouge_1_p = {}
    rouge_1_r = {}
    rouge_2_f = {}
    rouge_2_p = {}
    rouge_2_r = {}
    rouge_3_f = {}
    rouge_3_p = {}
    rouge_3_r = {}
    rouge_4_f = {}
    rouge_4_p = {}
    rouge_4_r = {}
    rouge_l_f = {}
    rouge_l_p = {}
    rouge_l_r = {}
    rouge_w_f = {}
    rouge_w_p = {}
    rouge_w_r = {}

    for i,path in enumerate(paths):
        files = os.listdir(path)

        for file_name in files:
            file = open(path + file_name, "r")
            lines = file.read().split('\n')
            if str(i) not in rouge_1_f:
                rouge_1_f[str(i)] = [float(lines[1][3:])]
            else:
                rouge_1_f[str(i)].append(float(lines[1][3:]))
            if str(i) not in rouge_1_p:
                rouge_1_p[str(i)] = [float(lines[2][3:])]
            else:
                rouge_1_p[str(i)].append(float(lines[2][3:]))
            if str(i) not in rouge_1_r:
                rouge_1_r[str(i)] = [float(lines[3][3:])]
            else:
                rouge_1_r[str(i)].append(float(lines[3][3:]))
            if str(i) not in rouge_2_f:
                rouge_2_f[str(i)] = [float(lines[5][3:])]
            else:
                rouge_2_f[str(i)].append(float(lines[5][3:]))
            if str(i) not in rouge_2_p:
                rouge_2_p[str(i)] = [float(lines[6][3:])]
            else:
                rouge_2_p[str(i)].append(float(lines[6][3:]))
            if str(i) not in rouge_2_r:
                rouge_2_r[str(i)] = [float(lines[7][3:])]
            else:
                rouge_2_r[str(i)].append(float(lines[7][3:]))
            if str(i) not in rouge_3_f:
                rouge_3_f[str(i)] = [float(lines[9][3:])]
            else:
                rouge_3_f[str(i)].append(float(lines[9][3:]))
            if str(i) not in rouge_3_p:
                rouge_3_p[str(i)] = [float(lines[10][3:])]
            else:
                rouge_3_p[str(i)].append(float(lines[10][3:]))
            if str(i) not in rouge_3_r:
                rouge_3_r[str(i)] = [float(lines[11][3:])]
            else:
                rouge_3_r[str(i)].append(float(lines[11][3:]))
            if str(i) not in rouge_4_f:
                rouge_4_f[str(i)] = [float(lines[13][3:])]
            else:
                rouge_4_f[str(i)].append(float(lines[13][3:]))
            if str(i) not in rouge_4_p:
                rouge_4_p[str(i)] = [float(lines[14][3:])]
            else:
                rouge_4_p[str(i)].append(float(lines[14][3:]))
            if str(i) not in rouge_4_r:
                rouge_4_r[str(i)] = [float(lines[15][3:])]
            else:
                rouge_4_r[str(i)].append(float(lines[15][3:]))
            if str(i) not in rouge_l_f:
                rouge_l_f[str(i)] = [float(lines[17][3:])]
            else:
                rouge_l_f[str(i)].append(float(lines[17][3:]))
            if str(i) not in rouge_l_p:
                rouge_l_p[str(i)] = [float(lines[18][3:])]
            else:
                rouge_l_p[str(i)].append(float(lines[18][3:]))
            if str(i) not in rouge_l_r:
                rouge_l_r[str(i)] = [float(lines[19][3:])]
            else:
                rouge_l_r[str(i)].append(float(lines[19][3:]))
            if str(i) not in rouge_w_f:
                rouge_w_f[str(i)] = [float(lines[21][3:])]
            else:
                rouge_w_f[str(i)].append(float(lines[21][3:]))
            if str(i) not in rouge_w_p:
                rouge_w_p[str(i)] = [float(lines[22][3:])]
            else:
                rouge_w_p[str(i)].append(float(lines[22][3:]))
            if str(i) not in rouge_w_r:
                rouge_w_r[str(i)] = [float(lines[23][3:])]
            else:
                rouge_w_r[str(i)].append(float(lines[23][3:]))


    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_f_rouge_1 = [np.array(rouge_1_f["0"]), np.array(rouge_1_f["1"]), np.array(rouge_1_f["2"]), np.array(rouge_1_f["3"]), np.array(rouge_1_f["4"])]
    bp_p = ax.boxplot(all_f_rouge_1, sym='')
    ax.set_title("f1 score 1_gram.")
    plt.savefig("f_rouge_1.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_p_rouge_1 = [np.array(rouge_1_p["0"]), np.array(rouge_1_p["1"]), np.array(rouge_1_p["2"]), np.array(rouge_1_p["3"]), np.array(rouge_1_p["4"])]
    bp_p = ax.boxplot(all_p_rouge_1, sym='')
    ax.set_title("precision 1_gram.")
    plt.savefig("p_rouge_1.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_r_rouge_1 = [np.array(rouge_1_r["0"]), np.array(rouge_1_r["1"]), np.array(rouge_1_r["2"]), np.array(rouge_1_r["3"]), np.array(rouge_1_r["4"])]
    bp_p = ax.boxplot(all_r_rouge_1, sym='')
    ax.set_title("recall 1_gram.")
    plt.savefig("r_rouge_1.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_f_rouge_2 = [np.array(rouge_2_f["0"]), np.array(rouge_2_f["1"]), np.array(rouge_2_f["2"]), np.array(rouge_2_f["3"]), np.array(rouge_2_f["4"])]
    bp_p = ax.boxplot(all_f_rouge_2, sym='')
    ax.set_title("f1 score 2_gram.")
    plt.savefig("f_rouge_2.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_p_rouge_2 = [np.array(rouge_2_p["0"]), np.array(rouge_2_p["1"]), np.array(rouge_2_p["2"]), np.array(rouge_2_p["3"]), np.array(rouge_2_p["4"])]
    bp_p = ax.boxplot(all_p_rouge_2, sym='')
    ax.set_title("precision 2_gram.")
    plt.savefig("p_rouge_2.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_r_rouge_2 = [np.array(rouge_2_r["0"]), np.array(rouge_2_r["1"]), np.array(rouge_2_r["2"]), np.array(rouge_2_r["3"]), np.array(rouge_2_r["4"])]
    bp_p = ax.boxplot(all_r_rouge_2, sym='')
    ax.set_title("recall 2_gram. ")
    plt.savefig("r_rouge_2.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_f_rouge_3 = [np.array(rouge_3_f["0"]), np.array(rouge_3_f["1"]), np.array(rouge_3_f["2"]), np.array(rouge_3_f["3"]), np.array(rouge_3_f["4"])]
    bp_p = ax.boxplot(all_f_rouge_3, sym='')
    ax.set_title("f1 score 3_gram.")
    plt.savefig("f_rouge_3.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_p_rouge_3 = [np.array(rouge_3_p["0"]), np.array(rouge_3_p["1"]), np.array(rouge_3_p["2"]), np.array(rouge_3_p["3"]), np.array(rouge_3_p["4"])]
    bp_p = ax.boxplot(all_p_rouge_3, sym='')
    ax.set_title("precision 3_gram.")
    plt.savefig("p_rouge_3.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_r_rouge_3 = [np.array(rouge_3_r["0"]), np.array(rouge_3_r["1"]), np.array(rouge_3_r["2"]), np.array(rouge_3_r["3"]), np.array(rouge_3_r["4"])]
    bp_p = ax.boxplot(all_r_rouge_3, sym='')
    ax.set_title("recall 3_gram.")
    plt.savefig("r_rouge_3.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_f_rouge_4 = [np.array(rouge_4_f["0"]), np.array(rouge_4_f["1"]), np.array(rouge_4_f["2"]), np.array(rouge_4_f["3"]), np.array(rouge_4_f["4"])]
    bp_p = ax.boxplot(all_f_rouge_4, sym='')
    ax.set_title("f1 score 4_gram.")
    plt.savefig("f_rouge_4.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_p_rouge_4 = [np.array(rouge_4_p["0"]), np.array(rouge_4_p["1"]), np.array(rouge_4_p["2"]), np.array(rouge_4_p["3"]), np.array(rouge_4_p["4"])]
    bp_p = ax.boxplot(all_p_rouge_4, sym='')
    ax.set_title("precision 4_gram.")
    plt.savefig("p_rouge_4.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_r_rouge_4 = [np.array(rouge_4_r["0"]), np.array(rouge_4_r["1"]), np.array(rouge_4_r["2"]), np.array(rouge_4_r["3"]), np.array(rouge_4_r["4"])]
    bp_p = ax.boxplot(all_r_rouge_4, sym='')
    ax.set_title("recall 4_gram. bert:")
    plt.savefig("r_rouge_4.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_f_rouge_l = [np.array(rouge_l_f["0"]), np.array(rouge_l_f["1"]), np.array(rouge_l_f["2"]), np.array(rouge_l_f["3"]), np.array(rouge_l_f["4"])]
    bp_p = ax.boxplot(all_f_rouge_l, sym='')
    ax.set_title("f1 score longest common sequence.")
    plt.savefig("f_rouge_l.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_p_rouge_l = [np.array(rouge_l_p["0"]), np.array(rouge_l_p["1"]), np.array(rouge_l_p["2"]), np.array(rouge_l_p["3"]), np.array(rouge_l_p["4"])]
    bp_p = ax.boxplot(all_p_rouge_l, sym='')
    ax.set_title("precision longest common sequence.")
    plt.savefig("p_rouge_l.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_r_rouge_l = [np.array(rouge_l_r["0"]), np.array(rouge_l_r["1"]), np.array(rouge_l_r["2"]), np.array(rouge_l_r["3"]), np.array(rouge_l_r["4"])]
    bp_p = ax.boxplot(all_r_rouge_l, sym='')
    ax.set_title("recall longest common sequence.")
    plt.savefig("r_rouge_l.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_f_rouge_w = [np.array(rouge_w_f["0"]), np.array(rouge_w_f["1"]), np.array(rouge_w_f["2"]), np.array(rouge_w_f["3"]), np.array(rouge_w_f["4"]) ]
    bp_p = ax.boxplot(all_f_rouge_w, sym='')
    ax.set_title("f1 score Weighted LCS-based. ")
    plt.savefig("f_rouge_w.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_p_rouge_w = [np.array(rouge_w_p["0"]), np.array(rouge_w_p["1"]), np.array(rouge_w_p["2"]), np.array(rouge_w_p["3"]), np.array(rouge_w_p["4"])]
    bp_p = ax.boxplot(all_p_rouge_w, sym='')
    ax.set_title("precision Weighted LCS-based.")
    plt.savefig("p_rouge_w.png")

    fig = plt.figure(figsize=(10, 7))
    ax = fig.add_axes([0.15, 0.1, 0.7, 0.8])
    ax.set_ylim([0, 1])
    ax.set_xticklabels(["bert", "bert2", "bert3", "bert4", "word2vec"])
    all_r_rouge_w = [np.array(rouge_w_r["0"]), np.array(rouge_w_r["1"]), np.array(rouge_w_r["2"]), np.array(rouge_w_r["3"]), np.array(rouge_w_r["4"])]
    bp_p = ax.boxplot(all_r_rouge_w, sym='')
    ax.set_title("recall Weighted LCS-based.")
    plt.savefig("r_rouge_w.png")









